"""文档内容抽取

优先使用 MinerU（结构化解析 PDF/DOCX 为 Markdown，保留标题层级），
失败时降级到轻量方案（PyMuPDF / python-docx / 明文）。
.txt/.md 直接读取。
.doc 旧格式暂不支持（需 LibreOffice 转换）。
"""
import logging
import subprocess
import tempfile
from pathlib import Path

from config import settings
from utils import mineru_api

logger = logging.getLogger("native_rag")

SUPPORTED_TYPES = {"pdf", "docx", "txt", "md"}

# MinerU 官方 API 单文件页数上限（官方硬限制），超限必失败，直接走本地 CLI 避免白等
MAX_API_PAGES = 200

# MinerU CLI 超时保险丝：
# 固定 900s 对数百页大 PDF 不够——CPU 逐页解析、扫描件更慢，会误触发降级丢结构化信息。
# 改为按页数动态估算，语义仍是"防永久卡死"而非"限时完成"。
MIN_TIMEOUT = 900          # 秒，基础预算（模型加载/预热 + 小文件）
PER_PAGE_BUDGET = 10       # 秒/页，覆盖文本型与扫描件余量


def _compute_cli_timeout(pages: int | None) -> int:
    """动态计算 MinerU CLI 超时：max(基础预算, 页数 × 每页预算)；非 PDF 用默认保险丝"""
    if pages:
        return max(MIN_TIMEOUT, pages * PER_PAGE_BUDGET)
    return MIN_TIMEOUT


def _page_count(path: str) -> int:
    """PyMuPDF 快查 PDF 页数（毫秒级，用于大 PDF 分流判断）"""
    import fitz

    with fitz.open(path) as doc:
        return doc.page_count


def _has_text_layer(path: str) -> bool:
    """采样判断 PDF 是否有文本层（文本型 vs 扫描件）

    采样前 2 + 中间 1 + 后 2 共 5 页，任一页 get_text() 非空即判文本型。
    混合型 PDF 丢少量扫描页可接受——总比整份走本地 MinerU（CPU 上 1 小时+）划算。
    """
    import fitz

    with fitz.open(path) as doc:
        total = doc.page_count
        indices = sorted({0, 1, total // 2, total - 2, total - 1})
        for i in indices:
            if 0 <= i < total and doc[i].get_text().strip():
                return True
    return False


def extract_text(file_path: str, file_type: str) -> str:
    """从文档文件抽取纯文本

    抽取路径（按文件类型分流，避免慢速本地 MinerU 大材小用）：
    - PDF：
      1. 配置了 Token 且页数 ≤ 200 → 官方 API（云端，快）
      2. 其余情况按文本层判断：
         文本型（有文本层）→ PyMuPDF 秒级提取；
         扫描件（无文本层）→ 本地 MinerU OCR（超时按页数动态放大，失败降级 PyMuPDF）
    - DOCX：python-docx 提取段落+表格（DOCX 无扫描概念，不需要 MinerU）
    - TXT/MD：明文读取
    """
    if file_type == "pdf":
        # 页数预检：API 分流与 CLI 超时预算共用
        pages = _page_count(file_path)
        if mineru_api.is_api_enabled() and pages <= MAX_API_PAGES:
            logger.info("[extract] %s 共 %s 页，走 MinerU 官方 API", Path(file_path).name, pages)
            try:
                return mineru_api.extract_with_api(file_path, settings.mineru_api_token)
            except Exception as e:
                # API 网络/限流失败降级本地：文本型秒级 PyMuPDF，扫描件走本地 OCR，不直接 failed
                logger.warning("[extract] MinerU 官方 API 失败，降级本地提取: %s", e)
        if _has_text_layer(file_path):
            # 文本型：PyMuPDF 秒级提取；标题层级损失换 1 小时+ 的本地 MinerU 不划算
            logger.info("[extract] %s 文本型 PDF，PyMuPDF 提取（%s 页）", Path(file_path).name, pages)
            return _extract_pdf(file_path)
        # 扫描件：必须 OCR，MinerU 不可替代；超时按页数放大
        timeout = _compute_cli_timeout(pages)
        logger.warning(
            "[extract] %s 扫描件（无文本层），走本地 MinerU OCR（%s 页，超时=%ss）",
            Path(file_path).name, pages, timeout,
        )
        try:
            return _extract_with_mineru(file_path, timeout)
        except Exception as e:
            # OCR 失败（超时/报错）降级 PyMuPDF：扫描件多半提取为空，由"内容为空"明确报错而非静默脏数据
            logger.warning("[extract] MinerU OCR 失败，降级 PyMuPDF: %s", e)
            return _extract_pdf(file_path)
    if file_type == "docx":
        # DOCX 是文本容器（无扫描概念），python-docx 直接取段落+表格，不需要 MinerU
        logger.info("[extract] %s 走 python-docx 提取", Path(file_path).name)
        return _extract_docx(file_path)
    if file_type in ("txt", "md"):
        return _extract_plain(file_path)
    if file_type == "doc":
        raise ValueError("暂不支持 .doc 旧格式，请先转换为 .docx 或 .pdf")
    raise ValueError(f"不支持的文件类型: {file_type}")


def _extract_with_mineru(file_path: str, timeout: int = MIN_TIMEOUT) -> str:
    """调用 MinerU CLI 抽取，返回结构化 Markdown 文本

    timeout：CLI 保险丝，默认 900s 基础预算；大 PDF 由调用方按页数动态放大
    """
    with tempfile.TemporaryDirectory() as out_dir:
        # pipeline backend 更通用；模型首次使用自动下载（已缓存到 /root/.cache）
        result = subprocess.run(
            ["mineru", "-p", file_path, "-o", out_dir, "-b", "pipeline", "-m", "auto"],
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            err = (result.stderr or "")[-500:]
            raise ValueError(f"MinerU 处理失败: {err}")
        md_files = list(Path(out_dir).rglob("*.md"))
        if not md_files:
            raise ValueError("MinerU 未生成 markdown 输出")
        return md_files[0].read_text(encoding="utf-8")


def _extract_pdf(path: str) -> str:
    """PDF 降级：PyMuPDF 逐页抽取"""
    import fitz

    doc = fitz.open(path)
    parts = []
    for page in doc:
        parts.append(page.get_text())
    doc.close()
    return "\n".join(parts)


def _extract_docx(path: str) -> str:
    """DOCX 降级：python-docx 抽取段落 + 表格"""
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain(path: str) -> str:
    """TXT/MD：按编码尝试读取（UTF-8 优先，失败回退 GBK）"""
    for enc in ("utf-8", "gbk"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
        except OSError as e:
            raise ValueError(f"读取文件失败: {e}")
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
